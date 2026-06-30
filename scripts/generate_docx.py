"""
generate_docx.py

用途：
將 UBD 跨域課程設計資料轉換成 Word 文件。

使用前準備：
pip install python-docx

注意：
此腳本為基礎版本，未綁定特定 Word 範本。
若要套用使用者提供的 Word 範本，需要再擴充模板讀取與欄位填入功能。
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_heading(document, text, level=1):
    """新增標題。"""
    document.add_heading(text, level=level)


def add_paragraph(document, text):
    """新增一般段落。"""
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.font.size = Pt(11)
    return paragraph


def add_key_value_table(document, data):
    """新增二欄式基本資料表。"""
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"

    header_cells = table.rows[0].cells
    header_cells[0].text = "項目"
    header_cells[1].text = "內容"

    for key, value in data.items():
        row_cells = table.add_row().cells
        row_cells[0].text = str(key)
        row_cells[1].text = str(value)

    return table


def add_list_table(document, headers, rows):
    """新增一般表格。"""
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"

    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        header_cells[index].text = header

    for row in rows:
        row_cells = table.add_row().cells
        for index, value in enumerate(row):
            row_cells[index].text = str(value)

    return table


def create_course_docx(course_data, output_path="ubd_course_design.docx"):
    """建立課程設計 Word 文件。"""

    document = Document()

    # 文件標題
    title = document.add_heading(course_data.get("course_name", "UBD 跨域課程設計"), level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 一、課程基本資料
    add_heading(document, "一、課程基本資料表", level=1)

    basic_info = {
        "課程名稱": course_data.get("course_name", ""),
        "適用年級": course_data.get("grade", ""),
        "總節數": course_data.get("periods", ""),
        "跨領域科目": course_data.get("subjects", ""),
        "SDGs / 議題": course_data.get("sdgs", ""),
        "課程大概念": course_data.get("big_idea", ""),
        "設計理念": course_data.get("design_rationale", "")
    }

    add_key_value_table(document, basic_info)

    # 二、課程架構分析
    add_heading(document, "二、課程架構分析", level=1)
    add_heading(document, "1. 大概念 Big Idea", level=2)
    add_paragraph(document, course_data.get("big_idea", ""))

    add_heading(document, "2. 真實情境脈絡", level=2)
    add_paragraph(document, course_data.get("real_context", ""))

    add_heading(document, "3. 跨領域整合邏輯", level=2)
    add_list_table(
        document,
        ["領域 / 科目", "在課程中的角色", "與其他領域的連結"],
        course_data.get("interdisciplinary_logic", [])
    )

    # 三、素養與學習重點
    add_heading(document, "三、素養與學習重點", level=1)

    add_heading(document, "1. 核心素養", level=2)
    add_list_table(
        document,
        ["核心素養方向", "對應說明"],
        course_data.get("competencies", [])
    )

    add_heading(document, "2. 學習表現", level=2)
    add_list_table(
        document,
        ["領域 / 科目", "學習表現", "對應說明"],
        course_data.get("learning_performance", [])
    )

    add_heading(document, "3. 學習內容", level=2)
    add_list_table(
        document,
        ["領域 / 科目", "學習內容", "對應說明"],
        course_data.get("learning_content", [])
    )

    # 四、K-U-D 學習目標
    add_heading(document, "四、K-U-D 學習目標", level=1)

    kud = course_data.get("kud", {})
    add_list_table(
        document,
        ["類別", "學習目標"],
        [
            ["K：Know 知道", kud.get("know", "")],
            ["U：Understand 理解", kud.get("understand", "")],
            ["D：Do 執行", kud.get("do", "")]
        ]
    )

    # 五、GRASPS 表現任務
    add_heading(document, "五、GRASPS 表現任務", level=1)

    grasps = course_data.get("grasps", {})
    add_list_table(
        document,
        ["GRASPS 元素", "說明"],
        [
            ["G：Goal 目標", grasps.get("goal", "")],
            ["R：Role 角色", grasps.get("role", "")],
            ["A：Audience 觀眾", grasps.get("audience", "")],
            ["S：Situation 情境", grasps.get("situation", "")],
            ["P：Product / Performance 產出或表現", grasps.get("product", "")],
            ["S：Standards 成功標準", grasps.get("standards", "")]
        ]
    )

    # 六、評量規準
    add_heading(document, "六、評量規準 Rubrics", level=1)
    add_list_table(
        document,
        ["評量向度", "4 優秀", "3 良好", "2 基礎", "1 待加強"],
        course_data.get("rubrics", [])
    )

    # 七、單元教學活動設計
    add_heading(document, "七、單元教學活動設計", level=1)
    add_list_table(
        document,
        ["節次", "時間", "教學流程", "教師指導語 / 關鍵提問", "學生活動", "評量方式 / 工具"],
        course_data.get("lesson_flow", [])
    )

    # 八、差異化與延伸建議
    add_heading(document, "八、差異化與延伸建議", level=1)

    differentiation = course_data.get("differentiation", {})

    add_heading(document, "1. 學習扶助", level=2)
    add_paragraph(document, differentiation.get("support", ""))

    add_heading(document, "2. 充實挑戰", level=2)
    add_paragraph(document, differentiation.get("extension", ""))

    # 儲存文件
    document.save(output_path)
    return output_path


if __name__ == "__main__":
    sample_course_data = {
        "course_name": "從細胞到城市：用數據設計永續創新系統",
        "grade": "高中二年級",
        "periods": "2 節課",
        "subjects": "國文、生物、數學",
        "sdgs": "SDG 9 產業創新與基礎建設",
        "big_idea": "系統的穩定與創新，來自各部分之間的有效連結、資料回饋與功能調整。",
        "design_rationale": "本課程引導學生從生物系統出發，連結城市基礎建設與產業創新，透過數學資料分析與國文表達能力，提出永續創新提案。",
        "real_context": "城市交通、能源與產業發展都像生物體一樣需要協調不同系統。",
        "interdisciplinary_logic": [
            ["國文", "建構提案論述", "將分析結果轉化為說服性文字"],
            ["生物", "理解系統概念", "提供城市系統的類比基礎"],
            ["數學", "資料判讀與分析", "支持問題分析與提案依據"]
        ],
        "competencies": [
            ["A2 系統思考與解決問題", "學生分析真實問題並提出解決策略"],
            ["B1 符號運用與溝通表達", "學生以文字、圖表與口語表達成果"],
            ["C2 人際關係與團隊合作", "學生透過小組合作完成提案"]
        ],
        "learning_performance": [
            ["國文", "待依正式課綱查核", "對應說服性表達與文本組織"],
            ["生物", "待依正式課綱查核", "對應系統、構造與功能概念"],
            ["數學", "待依正式課綱查核", "對應資料判讀與圖表分析"]
        ],
        "learning_content": [
            ["國文", "待依正式課綱查核", "論述文本與表達策略"],
            ["生物", "待依正式課綱查核", "系統、構造與功能"],
            ["數學", "待依正式課綱查核", "資料分析與圖表判讀"]
        ],
        "kud": {
            "know": "學生知道生物系統、SDG 9 與資料圖表的基本概念。",
            "understand": "學生理解自然系統與社會系統都需要協調運作與回饋調整。",
            "do": "學生能判讀資料、建立類比、提出方案並進行發表。"
        },
        "grasps": {
            "goal": "提出一項改善城市基礎建設或產業創新的永續提案。",
            "role": "城市創新顧問。",
            "audience": "市府代表、社區居民或青年創業團隊。",
            "situation": "某城市面臨交通、能源、水資源或產業轉型問題。",
            "product": "一頁式提案與 3 分鐘口頭發表。",
            "standards": "提案需包含問題描述、資料依據、生物系統類比、數學分析與說服性表達。"
        },
        "rubrics": [
            ["概念理解", "能精準說明並遷移應用", "能說明主要概念", "能說明部分概念", "概念混淆"],
            ["跨域整合", "自然整合三科觀點", "能整合兩科以上", "連結較薄弱", "只呈現單一科目"],
            ["問題解決", "提案具體可行且有資料支持", "提案合理", "方向可理解但支持不足", "提案模糊"],
            ["表達溝通", "清楚有說服力", "能清楚傳達重點", "組織略鬆散", "表達不清"],
            ["合作參與", "分工明確且協作有效", "多數成員參與", "參與不均", "缺乏合作"]
        ],
        "lesson_flow": [
            ["第一節", "0-10 分鐘", "引起動機", "如果一座城市是一個生命體，它的血管、神經與器官可能是什麼？", "學生自由聯想", "口頭回應"],
            ["第一節", "10-20 分鐘", "概念建構", "生物系統如何透過不同構造合作維持生命？", "閱讀資料並整理概念", "學習單"],
            ["第一節", "20-30 分鐘", "跨域連結", "城市系統哪裡像生物系統？", "完成類比表", "小組紀錄"],
            ["第一節", "30-50 分鐘", "資料判讀", "資料透露了什麼問題？", "判讀圖表", "圖表分析單"],
            ["第二節", "0-10 分鐘", "任務說明", "你們是城市創新顧問，要提出改善方案。", "理解任務", "任務檢核表"],
            ["第二節", "10-25 分鐘", "提案設計", "你的方案解決什麼問題？", "完成提案草稿", "教師回饋"],
            ["第二節", "25-40 分鐘", "發表回饋", "哪一組方案最兼顧創新與可行？", "小組發表", "同儕評量"],
            ["第二節", "40-50 分鐘", "反思收斂", "你如何用不同學科理解同一問題？", "完成個人反思", "出口票"]
        ],
        "differentiation": {
            "support": "提供系統對照表、資料判讀句型、一頁式提案模板與小組角色分工。",
            "extension": "加入第二組資料佐證，比較不同城市案例，並增加可行性與利害關係人分析。"
        }
    }

    create_course_docx(sample_course_data)